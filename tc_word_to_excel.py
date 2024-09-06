# -*- coding: utf-8 -*-
'''
@author: olivia.dou
Created on: 2023/6/24 8:34
desc:
使用步骤
1. 将现有Word用例去除其他部分(如封面页、文档描述等)，仅保留用例章节，单独保存一份文档（docx格式）
2. 修改tc_fields_loc中各字段对应值在Word用例表格中的位置(行,列),以0起始
3. 修改root_dir，source_path，dest_wb_path
4. tc_steps_prefix和expected_result_prefix（去除前缀后为需要提取的内容）
5. 运行脚本
6. 如果用例中有截图，将embeded_images_path对应路径下生成的图片批量上传到metersphere服务器/app/metersphere/data/image/markdown路径下（需要root权限）
7. 将Excel文件导入Metersphere平台


'''

import os,openpyxl,docx2txt,shutil
import traceback
from xml.dom import minidom
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl.drawing.image import Image
from msphere_lib import generate_random_string

# 指定测试用例中各字段值的位置(行,列), 索引从0开始
tc_fields_loc = {"ID":(1,1),"用例名称":(2,1),"所属模块":None,'前置条件':None,'备注':(3,1),'步骤描述':(4,0),'预期结果':(5,0),
                 '编辑模式':None,'标签':None,'责任人':(0,1),'用例等级':None,'用例状态':None,'用例类型':None,'附加字段':(6,0)}
tc_steps_prefix = "预置条件和测试步骤："
expected_result_prefix = "预期结果： \n"

current_file_path = os.path.abspath(__file__)
root_dir = os.path.dirname(current_file_path)
source_path = root_dir + os.sep + "existing_word_tc.docx"
source_doc = Document(source_path)
dest_wb_path = root_dir + os.sep +'dummy_output.xlsx' #目标用例workbook
dest_wb = openpyxl.Workbook()
dest_wb.save(dest_wb_path) # 同名覆盖，否则创建
#dest_sheet = dest_wb.create_sheet(dest_sheet_name)
dest_sheet = dest_wb.active
dest_wb.close()

#写入标题：
head = ['ID','用例名称','所属模块','前置条件','备注','步骤描述','预期结果','编辑模式','标签','责任人','用例等级','用例状态','用例类型','附加字段']
dest_sheet.append(head)

full_path_titles=[]
titles = []


def is_directory_empty(folder_path):
    if not os.path.isdir(folder_path):
        raise ValueError("Invalid folder path.")

    if len(os.listdir(folder_path)) == 0:
        return True
    else:
        return False


def clear_directory(folder_path):
    if not os.path.isdir(folder_path):
        raise ValueError("Invalid folder path.")

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)

# extract text and write images in Temporary Image directory
embeded_images_path = root_dir + os.sep + 'source_images/word/media'
if not os.path.exists(embeded_images_path):
    os.makedirs(embeded_images_path)
#如果文件夹非空，则清空文件夹
if not is_directory_empty(embeded_images_path):
    clear_directory(embeded_images_path )
docx2txt.process(source_path,embeded_images_path)

def get_case_module_path(full_path_titles, titles):
    """

    :param full_path_titles: 完整路径的标题列表
    :param titles: 当前标题列表

    :return: 用例所在模块
    """
    try:
        #在prev_titles的基础上用当前标题列表按照层级替换
        title_level = titles[0][0] #从titles的第一个level开始比对
        for i,title in enumerate(full_path_titles):
            if title[0]==title_level and title_level in [level for (level,title) in titles]:
                full_path_titles[i]=titles[[level for (level,title) in titles].index(title_level)]
                title_level+=1
            elif title[0]==title_level: #titles路径短
                for index in range(i,len(full_path_titles)):
                    del full_path_titles[index]

        if title_level in [level for (level,title) in titles]: #titles里有后续level的标题
            index = [level for (level,title) in titles].index(title_level)
            full_path_titles.extend(titles[index:])

        return full_path_titles
    except Exception as e:
        traceback.print_exc(e)


def iter_block_items(parent):
    # https://github.com/python-openxml/python-docx/issues/40
    from docx.document import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    # print('parent_elm: '+str(type(parent_elm)))
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)  # No recursion, return tables as tables
        # table = Table(child, parent)  # Use recursion to return tables as paragraphs
        # for row in table.rows:
        #     for cell in row.cells:
        #         yield from iter_block_items(cell)

def convert_content_to_list(content, prefix):
    """处理测试步骤、预期结果内容，按换行符拆分成列表

    :param content: 获取到的测试步骤单元格全部内容
    :param prefix: 需去除的前缀内容

    :return: 处理完的列表
    """
    content = content.replace(prefix, "").strip(' ')
    content_list = content.split('\n')
    return content_list


def insert_actual_result(source_doc, source_cell,target_sheet, target_row_index,target_col_index):
    """

    :param source_cell:
    :param target_row_index:
    :param target_col_index:

    :return:
    """

    for paragraph in source_cell.paragraphs:
        for run in paragraph.runs:
            #if run._r.xml.startswith('<w:drawing'):
            if '<w:drawing' in run._r.xml:
                # Parse the xml code for the blip
                xml_parser = minidom.parseString(run._r.xml)

                items = xml_parser.getElementsByTagName('a:blip')

                if items:
                    for i,item in enumerate(items):
                        image_path = item.getAttribute('r:embed')
                        # image_path = item._element.find('.//a:blip', run._r.nsmap).get(
                        #     '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        image_filename = source_doc.part.related_parts[image_path].partname.replace("/word/media/","")
                        print(image_filename)

                        # # 将图片插入到Excel表格单元格
                        # img = Image(embeded_images_path + image_filename)
                        # img.width = 100
                        # img.height = 100
                        # target_sheet.column_dimensions[
                        #     target_sheet.cell(row=target_row_index, column=target_col_index).column_letter].width = img.width / 6
                        # target_sheet.row_dimensions[target_row_index].height = img.height
                        # #target_sheet.cell(row=target_row_index, column=target_col_index).value = None
                        # img.anchor=get_column_letter(target_col_index) + str(target_row_index)
                        # target_sheet.add_image(img)

                        # 插入图片无法导入Metersphere，需插入图片路径
                        # 生成随机字符串，重命名文件（避免服务器上文件重名覆盖）
                        new_name = generate_random_string(8) + ".png"
                        os.chdir(embeded_images_path)
                        os.rename(image_filename,new_name)
                        target_sheet.cell(row=target_row_index, column=target_col_index).value += \
                            "![image.png](/resource/md/get?fileName=" + new_name + ')'

            else:
                target_sheet.cell(row=target_row_index, column=target_col_index).value += run.text


blocks = iter_block_items(source_doc)
last_case_row_index = 1
try:
    for block in blocks:
        #print(block.text if isinstance(block, Paragraph) else '<table>')
        if isinstance(block, Paragraph):
            print(block.text)
            if block.style.name is not None and block.style.name.startswith('Heading'):
                level = int(block.style.name[-1])  # 获取标题级别
                title = block.text  # 获取标题文本
                if level in [lev for (lev,tit) in titles]: # titles里已有该级别的标题
                    index = [lev for (lev,tit) in titles].index(level) # 获取该级别标题的索引位置
                    titles[index] = (level,title) # 将该级别的标题改为当前标题
                    len_titles = len(titles)
                    i=index+1
                    while i<len(titles): # 当前级别后还有元素
                        del titles[i] # 删除后面的元素
                else:
                    titles.append((level, title))
                #print(titles)
        elif isinstance(block, Table):
            # for row in block.rows:
            #     row_data = []
            #     for cell in row.cells:
            #         for paragraph in cell.paragraphs:
            #             row_data.append(paragraph.text)
            #     print("\t".join(row_data))
            table = block
            # tables.append(table)
            # 获取用例对应的模块完整路径
            full_path_titles = get_case_module_path(full_path_titles, titles)
            # 将标题处理成模块
            module = "/" + "/".join([title for i,title in full_path_titles])
            print(module)
            # 读取用例信息
            tc_steps = convert_content_to_list(table.cell(*tc_fields_loc['步骤描述']).text, tc_steps_prefix)
            expected_results =convert_content_to_list(table.cell(*tc_fields_loc['预期结果']).text, expected_result_prefix)
            n_rows = max(len(tc_steps),len(expected_results))
            tc_data = []

            for i in range(min(len(tc_steps),len(expected_results))):
                tc_data.append({})
                tc_data[i]['tc_step'] = tc_steps[i]
                tc_data[i]['expected_result'] = expected_results[i]

            for i in range(min(len(tc_steps),len(expected_results)),n_rows): #两者长度不等
                tc_data.append({})
                if len(tc_steps)<len(expected_results):
                    tc_data[i]['tc_step']=''
                    tc_data[i]['expected_result'] = expected_results[i]
                else:
                    tc_data[i]['tc_step'] = tc_steps[i]
                    tc_data[i]['expected_result'] = ''


            tc_id = table.cell(*tc_fields_loc['ID']).text
            tc_name = table.cell(*tc_fields_loc['用例名称']).text
            module = module
            precondition = '' #视情况而定，合肥项目模板中前置条件和测试步骤混在一起无法区分，因此放弃填充前置条件，统一作为测试步骤处理
            memo= table.cell(*tc_fields_loc['备注']).text
            edit_mode = 'STEP'
            status = '已完成'
            owner = table.cell(*tc_fields_loc['责任人']).text
            priority = 'P2'
            case_type = '正常用例'
            label = ''
            additional_field = ''

            for i in range(n_rows):
                if i==0:
                    dest_sheet.append([tc_id,tc_name,module,precondition,memo,tc_data[i]['tc_step'],
                                      tc_data[i]['expected_result'],edit_mode,label,owner,priority,status,case_type,additional_field])
                    #将实际结果插入附加字段
                    source_cell = table.cell(*tc_fields_loc['附加字段'])
                    #print(source_cell._tc.xml)
                    insert_actual_result(source_doc,source_cell,dest_sheet,last_case_row_index+1,head.index('附加字段')+1)
                else:
                    dest_sheet.append(['', '', '',  '', '', tc_data[i]['tc_step'], tc_data[i]['expected_result'],
                                       '', '', '', '', '', '', ''])


            # 合并单元格
            if n_rows>1:
                for col_index in [head.index('ID'),head.index('用例名称'),head.index('所属模块'),head.index('前置条件'),
                        head.index('备注'),head.index('编辑模式'),head.index('标签'),head.index('责任人'),head.index('用例等级'),
                                  head.index('用例状态'),head.index('用例类型'),head.index('附加字段')]:
                    dest_sheet.merge_cells(start_row=last_case_row_index + 1, end_row=last_case_row_index+n_rows, start_column=col_index+1,
                                           end_column=col_index+1)
            last_case_row_index = last_case_row_index+n_rows


    dest_wb.save(dest_wb_path)

except Exception as e:
    traceback.print_exc()



