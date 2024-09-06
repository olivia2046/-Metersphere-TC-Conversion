# -*- coding: utf-8 -*-
'''
@author: olivia.dou
Created on: 2023/5/18 15:07
desc:
使用方法：
如果导出用例需要附带截图，则本脚本需在Metersphere所在服务器执行（否则找不到图片）
如果不需要附带截图，则脚本可在本地执行
1.Metersphere中按顺序（如用例ID）排列全部用例，全部导出，勾选全部基础字段、自定义字段，以及其他字段中的评论、创建人、创建时间（排序用）、更新时间等
2.检查导出的Excel用例的顺序，如果和预期输出的用例顺序不一致，手动调整（也可在生成WORD后调整WORD文档中用例顺序）
3.修改path （项目模板文件目录下需创建好用例模板文件，包含项目WORD版用例中的章节号和表格，表格内填入邮件合并域）
4.修改src_wb_path
5.修改actual_result_row_index，actual_result_col_index（"实际结果"字段内容在表格中的位置索引，以0起始）
6.设置need_image(若输出的用例不需要带截图，则设置为False)。需要输出截图时注意设置img_width和img_height
7.修改image_dir(用例截图所在路径，服务器上为固定地址，仅调试脚本时需要修改)
8.使用python3运行脚本
注：插入图片域一个域只能替换一张图片，由于一个用例可能需要在实际结果中插入多张图片，因此在完成邮件合并后使用文件遍历方式插入实际结果字段的内容
插入图片域参考：https://stackoverflow.com/questions/54260657/python-script-insert-image-using-mailmerge-into-docx-file
'''
import os, traceback

from docx.shared import Inches
from openpyxl import open
from docx import Document
from mailmerge import MailMerge # install docx-mailmerge
from msphere_lib import get_parts_from_comment_field,get_tc_data

#n_levels = 3


path = "."
template_path = path + os.sep + "Testcase_template.docx"
template_doc = Document(template_path)
src_wb_path = path + os.sep + "Metersphere_case_dummy.xlsx"
output_path = path + os.sep + "Word_tc_output.docx"
#实际结果在模板中的位置
actual_result_row_index = 6
actual_result_col_index = 0
need_image=False #输出的用例是否需要带截图
# PC端可设置4*2，移动端可设置2*3
img_width = Inches(4)
img_height = Inches(2)
#image_dir = path + os.sep + 'images'
image_dir = r"/app/metersphere/data/image/markdown"
print("图片路径：%s"%image_dir)
print("读取Excel用例文件")
src_wb = open(src_wb_path)
src_sheet_name = "模版" #注意Metersphere导出的Excel文件sheet名为模版，而非模板
src_sheet = src_wb[src_sheet_name]

# data = pd.read_excel(src_wb_path,src_sheet_name)
#columns = ['ID','用例名称', '所属模块', '前置条件', '备注', '步骤描述', '预期结果', '编辑模式', '标签', '责任人', '用例等级', '用例状态', '用例类型', '附加字段']
head_row = src_sheet[1]
columns = [cell.value for cell in head_row if cell.value is not None]
# data[columns] = data[columns].fillna(method='ffill')
# print(data.values.tolist())


tc_data = []
tc_row = {}
num_rows = 0
previous_module_paths = []
print("读取用例信息")
try:

    tc_data = get_tc_data(src_wb_path, src_sheet_name)
    print("开始邮件合并...")
    merge_doc = MailMerge(template_path)
    merge_doc.merge_templates(tc_data, separator='textWrapping_break')
    merge_doc.write(output_path)
    #fill_table_data(template_path, output_path, tc_data)
    print("邮件合并结束")

    #扫描一遍文档，None的段落删除
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    output_doc = Document(output_path)
    for para in output_doc.paragraphs:
        #print('段落文本：%s'%para.text)
        if para.text=='':
            delete_paragraph(para)

    output_doc.save(output_path)
    #插入实际结果
    def insert_actual_result(paragraph, actual_result, image_dir):
        results = get_parts_from_comment_field(actual_result,image_dir)

        for result in results:
            run = paragraph.add_run()
            run.add_text(result[0])
            image_path = result[1]

            if len(image_path) > 0 and need_image:
                run = paragraph.add_run()
                run.add_picture(image_path, width=img_width, height=img_height)
            run.add_text(result[2])


    for i,table in enumerate(output_doc.tables):
        para = table.cell(6, 0).add_paragraph()
        insert_actual_result(para,tc_data[i]['actual_result'],image_dir)

    output_doc.save(output_path)


except Exception as e:
    traceback.print_exc()




